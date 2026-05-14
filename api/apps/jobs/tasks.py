import logging
import httpx
import json
import os
import re
import tempfile
import fitz  # PyMuPDF
import docx
from celery import shared_task
from django.core.cache import cache
from decouple import config
from .models import JobPostActivity

logger = logging.getLogger(__name__)


def _truncate_text(value: str, max_chars: int) -> str:
    text = (value or "").strip()
    if max_chars <= 0 or len(text) <= max_chars:
        return text

    truncated = text[:max_chars].rsplit(" ", 1)[0].strip()
    return f"{truncated}\n...[truncated]"


def _acquire_analysis_slot(activity_id: int, max_slots: int, ttl_seconds: int) -> str | None:
    if max_slots <= 0:
        return None

    for idx in range(max_slots):
        slot_key = f"ai:resume-analysis:slot:{idx}"
        if cache.add(slot_key, str(activity_id), timeout=ttl_seconds):
            return slot_key

    return None


def _release_analysis_slot(slot_key: str | None):
    if slot_key:
        cache.delete(slot_key)


def _strip_code_fences(text: str) -> str:
    value = (text or "").strip()
    if "```json" in value:
        value = value.split("```json", 1)[1].split("```", 1)[0].strip()
    elif "```" in value:
        value = value.split("```", 1)[1].split("```", 1)[0].strip()
    return value


def _strip_reasoning_blocks(text: str) -> str:
    value = (text or "").strip()
    value = value.replace("\ufeff", "").replace("\u200b", " ")
    value = re.sub(r"<think>[\s\S]*?</think>", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"<think>[\s\S]*", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"</think>", " ", value, flags=re.IGNORECASE)
    return value.strip()


def _json_object_candidates(text: str):
    """Yield balanced JSON-object substrings from model output."""
    value = text or ""
    for start, char in enumerate(value):
        if char != "{":
            continue

        depth = 0
        in_string = False
        escape = False
        for idx in range(start, len(value)):
            current = value[idx]
            if in_string:
                if escape:
                    escape = False
                elif current == "\\":
                    escape = True
                elif current == '"':
                    in_string = False
                continue

            if current == '"':
                in_string = True
            elif current == "{":
                depth += 1
            elif current == "}":
                depth -= 1
                if depth == 0:
                    yield value[start:idx + 1]
                    break


def _parse_llm_json_content(raw_text: str) -> dict:
    cleaned = _strip_reasoning_blocks(_strip_code_fences(raw_text))
    if not cleaned:
        raise ValueError("Empty model content")

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    decoder = json.JSONDecoder()
    try:
        parsed, _ = decoder.raw_decode(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    for candidate in _json_object_candidates(cleaned):
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            continue

    raise ValueError("Model response does not contain valid JSON.")


def _coerce_list(value) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()][:8]
    if isinstance(value, str):
        parts = [part.strip(" -•\t\r\n") for part in value.split(",")]
        return [part for part in parts if part][:8]
    return []


def _normalize_analysis_result(result: dict) -> dict:
    score = result.get("score", result.get("overall_score", 0))
    try:
        score = int(float(score))
    except (TypeError, ValueError):
        score = 0
    score = max(0, min(100, score))

    return {
        "score": score,
        "summary": str(result.get("summary") or result.get("recommendation") or "").strip(),
        "skills": _coerce_list(result.get("skills")),
        "pros": _coerce_list(result.get("pros") or result.get("strengths")),
        "cons": _coerce_list(result.get("cons") or result.get("gaps")),
        "matching_skills": _coerce_list(result.get("matching_skills")),
        "missing_skills": _coerce_list(result.get("missing_skills")),
    }


@shared_task(
    bind=True,
    autoretry_for=(Exception,),
    retry_backoff=True,
    retry_kwargs={"max_retries": 3},
    ignore_result=True,
)
def es_index_job_post(self, job_post_id: int):
    """
    Asynchronously index (create or update) a single JobPost document in Elasticsearch.
    Called by the post_save signal in signals.py.
    """
    try:
        from .models import JobPost
        from django_elasticsearch_dsl.registries import registry

        instance = JobPost.objects.get(pk=job_post_id)
        registry.update(instance)
        logger.debug("ES: indexed JobPost id=%s", job_post_id)
    except JobPost.DoesNotExist:
        logger.warning("ES index skipped: JobPost id=%s not found (deleted?)", job_post_id)
    except Exception as exc:
        logger.error("ES index failed for JobPost id=%s: %s", job_post_id, exc)
        raise  # let Celery retry


@shared_task(
    bind=True,
    autoretry_for=(Exception,),
    retry_backoff=True,
    retry_kwargs={"max_retries": 3},
    ignore_result=True,
)
def es_delete_job_post(self, job_post_id: int):
    """
    Asynchronously remove a JobPost document from Elasticsearch.
    Called by the post_delete signal in signals.py.
    """
    try:
        from django_elasticsearch_dsl.registries import registry
        from .documents import JobPostDocument

        # Build a fake object with just the pk so the registry can remove it
        class _Stub:
            pk = job_post_id

        JobPostDocument().delete(doc_id=job_post_id, ignore=404)
        logger.debug("ES: deleted JobPost id=%s", job_post_id)
    except Exception as exc:
        logger.error("ES delete failed for JobPost id=%s: %s", job_post_id, exc)
        raise


def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
    return text

@shared_task(bind=True, autoretry_for=(httpx.TimeoutException, httpx.ConnectError),
             retry_backoff=True, retry_kwargs={'max_retries': 3})
def analyze_resume_ai(self, activity_id):
    """
    Task phan tich CV bang AI.
    Tu dong retry toi da 3 lan voi exponential backoff khi LLM timeout/connect error.
    """
    temp_file = None
    slot_key = None

    try:
        max_slots = config("AI_RESUME_ANALYSIS_MAX_CONCURRENCY", default=2, cast=int)
        slot_wait_seconds = config("AI_RESUME_ANALYSIS_SLOT_WAIT_SECONDS", default=20, cast=int)
        slot_ttl_seconds = config("AI_RESUME_ANALYSIS_SLOT_TTL_SECONDS", default=1800, cast=int)

        slot_key = _acquire_analysis_slot(
            activity_id=activity_id,
            max_slots=max_slots,
            ttl_seconds=slot_ttl_seconds,
        )
        if max_slots > 0 and not slot_key:
            logger.info(
                "Activity %s waiting AI slot (limit=%s), retry in %ss",
                activity_id,
                max_slots,
                slot_wait_seconds,
            )
            raise self.retry(countdown=slot_wait_seconds)

        activity = JobPostActivity.objects.select_related('job_post', 'resume', 'resume__file').get(id=activity_id)

        if not activity.resume or not activity.resume.file:
            activity.ai_analysis_status = 'failed'
            activity.ai_analysis_progress = 0
            activity.ai_analysis_summary = "Khong tim thay file CV de phan tich."
            activity.save(update_fields=['ai_analysis_status', 'ai_analysis_progress', 'ai_analysis_summary', 'update_at'])
            return

        activity.ai_analysis_status = 'processing'
        activity.ai_analysis_progress = 5
        activity.save(update_fields=['ai_analysis_status', 'ai_analysis_progress', 'update_at'])

        file_obj = activity.resume.file
        file_format = file_obj.format.lower() if file_obj.format else 'pdf'
        resume_text = ""

        # Strategy 1: direct MinIO fetch inside Docker network
        try:
            from shared.helpers.cloudinary_service import CloudinaryService
            from django.conf import settings as django_settings

            minio_client = CloudinaryService._get_client()
            bucket = django_settings.MINIO_BUCKET
            object_name = file_obj.public_id.lstrip('/')

            if object_name.startswith("http://") or object_name.startswith("https://"):
                from urllib.parse import urlparse
                parsed = urlparse(object_name)
                object_name = parsed.path.lstrip("/")
                if object_name.startswith(f"{bucket}/"):
                    object_name = object_name[len(bucket) + 1:]

            logger.info("Downloading CV from MinIO: bucket=%s, object=%s", bucket, object_name)

            response = minio_client.get_object(bucket, object_name)
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_format}") as tf:
                for chunk in response.stream(1024 * 32):
                    tf.write(chunk)
                temp_file = tf.name
            response.close()
            response.release_conn()

            activity.ai_analysis_progress = 25
            activity.save(update_fields=['ai_analysis_progress', 'update_at'])

        except Exception as minio_err:
            logger.warning("MinIO direct download failed: %s. Falling back to HTTP URL.", minio_err)

            resume_url = file_obj.get_full_url()
            logger.info("Downloading CV from URL: %s", resume_url)
            with httpx.Client(timeout=30.0) as client:
                response = client.get(resume_url)
                if response.status_code != 200:
                    raise Exception(f"Failed to download resume file: HTTP {response.status_code}")

                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_format}") as tf:
                    tf.write(response.content)
                    temp_file = tf.name

            activity.ai_analysis_progress = 25
            activity.save(update_fields=['ai_analysis_progress', 'update_at'])

        # Extract text from downloaded file
        if temp_file:
            if file_format == 'pdf':
                resume_text = extract_text_from_pdf(temp_file)
            elif file_format in ['docx', 'doc']:
                resume_text = extract_text_from_docx(temp_file)
            else:
                with open(temp_file, 'r', errors='ignore') as f:
                    resume_text = f.read()

            activity.ai_analysis_progress = 45
            activity.save(update_fields=['ai_analysis_progress', 'update_at'])

        if not resume_text or len(resume_text.strip()) < 50:
            activity.ai_analysis_status = 'failed'
            activity.ai_analysis_progress = 0
            activity.ai_analysis_summary = "Khong the doc duoc noi dung CV (co the la file anh hoac file loi)."
            activity.save(update_fields=['ai_analysis_status', 'ai_analysis_progress', 'ai_analysis_summary', 'update_at'])
            return

        # Compact prompt to reduce token use and improve throughput
        max_resume_chars = config("AI_RESUME_PROMPT_MAX_CV_CHARS", default=2600, cast=int)
        max_jd_chars = config("AI_RESUME_PROMPT_MAX_JD_CHARS", default=900, cast=int)
        max_req_chars = config("AI_RESUME_PROMPT_MAX_REQUIREMENT_CHARS", default=700, cast=int)

        job_description = _truncate_text(activity.job_post.job_description or "", max_jd_chars)
        job_requirement = _truncate_text(activity.job_post.job_requirement or "", max_req_chars)
        resume_excerpt = _truncate_text(resume_text, max_resume_chars)

        prompt = f"""
        Analyze candidate CV fit for this role and return ONLY one compact valid JSON object.

        Job title: {activity.job_post.job_name}
        Job description: {job_description}
        Job requirements: {job_requirement}

        Candidate CV:
        {resume_excerpt}

        Required JSON schema:
        {{
          "score": 0-100 integer,
          "summary": "short Vietnamese summary",
          "skills": ["..."],
          "pros": ["..."],
          "cons": ["..."],
          "matching_skills": ["..."],
          "missing_skills": ["..."]
        }}

        Constraints:
        - summary max 70 words.
        - each list max 5 short items.
        - no markdown, no extra text outside JSON.
        - no <think> block, no explanation.
        """

        llm_base_url = config(
            "AI_LLM_BASE_URL",
            default=config("LLM_BASE_URL", default=config("OLLAMA_BASE_URL", default="http://ollama:11434/v1")),
        )
        model_alias = config(
            "AI_RESUME_LLM_MODEL",
            default=config("AI_LLM_MODEL", default=config("LLM_MODEL", default=config("OLLAMA_MODEL", default="gemma4:e4b"))),
        )
        llm_temperature = config("AI_RESUME_LLM_TEMPERATURE", default=0.1, cast=float)
        llm_top_p = config("AI_RESUME_LLM_TOP_P", default=0.9, cast=float)
        llm_max_tokens = config("AI_RESUME_LLM_MAX_TOKENS", default=900, cast=int)
        llm_timeout = config("AI_RESUME_LLM_TIMEOUT_SECONDS", default=240.0, cast=float)
        llm_connect_timeout = config("AI_RESUME_LLM_CONNECT_TIMEOUT_SECONDS", default=10.0, cast=float)

        payload = {
            "model": model_alias,
            "messages": [
                {"role": "system", "content": "Ban la AI tuyen dung. Chi tra ve dung mot JSON object hop le, khong markdown, khong giai thich, khong <think>."},
                {"role": "user", "content": prompt},
            ],
            "temperature": llm_temperature,
            "top_p": llm_top_p,
            "max_tokens": llm_max_tokens,
            "stream": False,
            "response_format": {"type": "json_object"},
        }

        activity.ai_analysis_progress = 70
        activity.save(update_fields=['ai_analysis_progress', 'update_at'])

        llm_api_key = config("AI_LLM_API_KEY", default="")
        headers = {}
        if llm_api_key:
            headers["Authorization"] = f"Bearer {llm_api_key}"

        with httpx.Client(timeout=httpx.Timeout(timeout=llm_timeout, connect=llm_connect_timeout)) as client:
            resp = client.post(f"{llm_base_url}/chat/completions", json=payload, headers=headers)

        if resp.status_code != 200:
            raise Exception(f"LLM API failed with status {resp.status_code}")

        response_json = resp.json()
        message = (response_json.get("choices") or [{}])[0].get("message") or {}
        content = message.get("content") or ""
        reasoning = message.get("reasoning") or ""
        if isinstance(content, list):
            content = "\n".join(
                part.get("text", "") if isinstance(part, dict) else str(part)
                for part in content
            )

        # Ollama + some reasoning models may return empty "content" on /v1/chat/completions.
        # Fallback to native /api/chat to force JSON output when needed.
        try:
            result = _parse_llm_json_content(content)
        except Exception:
            result = None
            if reasoning:
                try:
                    result = _parse_llm_json_content(reasoning)
                except Exception:
                    result = None

            if result is None and config("AI_RESUME_OLLAMA_FALLBACK_ENABLED", default=True, cast=bool):
                native_base = llm_base_url[:-3] if llm_base_url.endswith("/v1") else llm_base_url
                native_payload = {
                    "model": model_alias,
                    "messages": payload["messages"],
                    "stream": False,
                    "think": False,
                    "format": "json",
                    "options": {
                        "temperature": llm_temperature,
                    },
                }
                with httpx.Client(timeout=httpx.Timeout(timeout=llm_timeout, connect=llm_connect_timeout)) as native_client:
                    native_resp = native_client.post(f"{native_base}/api/chat", json=native_payload, headers=headers)
                if native_resp.status_code == 200:
                    native_json = native_resp.json()
                    native_content = (native_json.get("message") or {}).get("content") or ""
                    result = _parse_llm_json_content(native_content)

            if result is None:
                raise ValueError("Model response does not contain valid JSON.")

        result = _normalize_analysis_result(result)
        activity.ai_analysis_score = result.get('score', 0)
        activity.ai_analysis_summary = result.get('summary', '')
        activity.ai_analysis_skills = ", ".join(result.get('skills', []))
        activity.ai_analysis_pros = ", ".join(result.get('pros', []))
        activity.ai_analysis_cons = ", ".join(result.get('cons', []))
        activity.ai_analysis_matching_skills = result.get('matching_skills', [])
        activity.ai_analysis_missing_skills = result.get('missing_skills', [])
        activity.ai_analysis_status = 'completed'
        activity.ai_analysis_progress = 100
        activity.save()

        logger.info("AI Analysis completed for Activity %s. Score: %s", activity_id, activity.ai_analysis_score)

    except (httpx.TimeoutException, httpx.ConnectError) as exc:
        max_r = 3  # matches retry_kwargs['max_retries']
        if (self.request.retries or 0) >= max_r:
            try:
                act = JobPostActivity.objects.get(id=activity_id)
                act.ai_analysis_status = 'failed'
                act.ai_analysis_progress = 0
                act.ai_analysis_summary = f"LLM khong phan hoi sau nhieu lan thu: {str(exc)[:300]}"
                act.save(update_fields=['ai_analysis_status', 'ai_analysis_progress', 'ai_analysis_summary', 'update_at'])
            except Exception:
                logger.error("Failed to update activity %s after final retry", activity_id)
        raise

    except Exception as e:
        logger.error("Error in analyze_resume_ai for activity %s: %s", activity_id, e)
        try:
            activity = JobPostActivity.objects.get(id=activity_id)
            activity.ai_analysis_status = 'failed'
            activity.ai_analysis_progress = 0
            activity.ai_analysis_summary = str(e)[:500]
            activity.save(update_fields=['ai_analysis_status', 'ai_analysis_progress', 'ai_analysis_summary', 'update_at'])
        except Exception:
            logger.error("Failed to update activity %s status after error", activity_id)

    finally:
        _release_analysis_slot(slot_key)

        if temp_file and os.path.exists(temp_file):
            try:
                os.unlink(temp_file)
            except OSError:
                pass

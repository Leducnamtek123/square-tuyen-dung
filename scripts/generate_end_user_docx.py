from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "docs-end-user"
IMG_DIR = DOC_DIR / "images"
DOCX_PATH = DOC_DIR / "HUONG_DAN_NGUOI_DUNG.docx"


def apply_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup_document(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)


def add_paragraph(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False, size: int = 12, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    apply_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        apply_font(run)


def add_numbers(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        apply_font(run)


def add_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Khu vực"
    hdr[1].text = "Mục đích"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def add_image(doc: Document, filename: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(IMG_DIR / filename), width=Inches(6.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    apply_font(run, size=10, italic=True)


def build() -> Document:
    doc = Document()
    setup_document(doc)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HƯỚNG DẪN SỬ DỤNG SQUARE TUYỂN DỤNG")
    apply_font(run, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Dành cho người dùng cuối, ứng viên và nhà tuyển dụng")
    apply_font(run, size=11, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Phiên bản 1.0 - 09/05/2026")
    apply_font(run, size=10)

    doc.add_paragraph("")

    doc.add_heading("1. Mục đích tài liệu", level=1)
    add_paragraph(
        doc,
        "Tài liệu này hướng dẫn cách sử dụng Square Tuyển Dụng từ góc nhìn người dùng cuối. Phần nội dung tập trung vào các thao tác công khai, thao tác tài khoản và các bước thường gặp khi tìm việc, xem công ty và ứng tuyển."
    )
    add_bullets(doc, [
        "Truy cập trang chủ và làm quen với thanh điều hướng.",
        "Đăng ký tài khoản, đăng nhập và hoàn thiện hồ sơ.",
        "Tìm kiếm việc làm, xem chi tiết tin và nộp hồ sơ.",
        "Tìm kiếm công ty, mở hồ sơ doanh nghiệp và theo dõi việc làm đang tuyển.",
        "Xử lý một số tình huống cơ bản như quên mật khẩu, không thấy kết quả hoặc không vào được màn hình."
    ])

    doc.add_heading("2. Tổng quan các khu vực", level=1)
    add_table(doc, [
        ("Trang chủ", "Giới thiệu nền tảng, đi tới các khu vực chính"),
        ("Việc làm", "Tìm việc theo từ khóa, ngành nghề và địa điểm"),
        ("Công ty", "Xem hồ sơ doanh nghiệp và vị trí đang tuyển"),
        ("Đăng nhập", "Vào hệ thống bằng tài khoản đã có"),
        ("Đăng ký", "Tạo tài khoản ứng viên mới"),
    ])

    doc.add_heading("3. Bắt đầu sử dụng", level=1)
    doc.add_heading("3.1 Truy cập trang chủ", level=2)
    add_paragraph(doc, "Khi mở trang chủ, bạn sẽ thấy:")
    add_bullets(doc, [
        "Thanh điều hướng ở phía trên cùng.",
        "Khối tìm kiếm việc làm ở phần đầu trang.",
        "Khối giới thiệu các công ty và tin tuyển dụng nổi bật.",
        "Lối đi nhanh dành cho ứng viên và nhà tuyển dụng.",
    ])
    add_image(doc, "01-home-top.png", "Hình 1. Trang chủ Square Tuyển Dụng")

    doc.add_heading("3.2 Đăng nhập", level=2)
    add_numbers(doc, [
        "Chọn Đăng nhập ở góc phải trên cùng.",
        "Chọn kiểu đăng nhập bằng Email hoặc SĐT nếu hệ thống cho phép.",
        "Nhập email hoặc số điện thoại.",
        "Nhập mật khẩu.",
        "Nhấn Đăng nhập để vào hệ thống.",
    ])
    add_paragraph(doc, "Nếu quên mật khẩu, chọn Quên mật khẩu? để khởi tạo lại.")
    add_image(doc, "06-login.png", "Hình 2. Màn hình đăng nhập")

    doc.add_heading("3.3 Đăng ký tài khoản mới", level=2)
    add_numbers(doc, [
        "Chọn Đăng ký.",
        "Nhập họ và tên.",
        "Nhập email.",
        "Tạo mật khẩu và xác nhận mật khẩu.",
        "Nhấn Đăng ký.",
    ])
    add_paragraph(doc, "Sau khi đăng ký, hệ thống có thể yêu cầu xác thực email trước khi cho phép dùng đầy đủ chức năng.")
    add_image(doc, "07-register.png", "Hình 3. Màn hình đăng ký tài khoản")

    doc.add_heading("4. Tìm kiếm việc làm", level=1)
    doc.add_heading("4.1 Dùng trang việc làm", level=2)
    add_paragraph(doc, "Trang Việc làm là nơi bạn tìm việc theo nhiều tiêu chí khác nhau.")
    add_bullets(doc, [
        "Từ khóa công việc.",
        "Ngành nghề.",
        "Tỉnh, thành phố.",
        "Bộ lọc nâng cao nếu muốn thu hẹp kết quả.",
    ])
    add_image(doc, "02-jobs-top.png", "Hình 4. Trang danh sách việc làm")

    doc.add_heading("4.2 Cách tìm nhanh một công việc phù hợp", level=2)
    add_numbers(doc, [
        "Nhập từ khóa, ví dụ: kỹ sư, thiết kế, kế toán.",
        "Chọn ngành nghề nếu cần lọc sâu hơn.",
        "Chọn tỉnh hoặc thành phố bạn muốn làm việc.",
        "Nhấn Tìm kiếm.",
        "Đọc danh sách kết quả và mở tin phù hợp nhất.",
    ])

    doc.add_heading("4.3 Ý nghĩa thông tin trên thẻ việc làm", level=2)
    add_bullets(doc, [
        "Tên công việc.",
        "Mức lương dự kiến.",
        "Địa điểm làm việc.",
        "Hạn nộp hồ sơ.",
        "Nhãn HOT hoặc Gấp nếu tin nổi bật.",
    ])

    doc.add_heading("4.4 Xem chi tiết tin tuyển dụng", level=2)
    add_numbers(doc, [
        "Chọn vào một tin việc làm.",
        "Xem phần tiêu đề, mức lương, hạn nộp và mô tả công việc.",
        "Đọc kỹ mục yêu cầu, quyền lợi, thông tin liên hệ và vị trí làm việc.",
        "Kiểm tra mục việc làm tương tự nếu muốn so sánh thêm.",
    ])
    add_image(doc, "03-job-detail-top.png", "Hình 5. Trang chi tiết việc làm")

    doc.add_heading("4.5 Lưu tin và ứng tuyển", level=2)
    add_numbers(doc, [
        "Chọn Nộp hồ sơ nếu muốn ứng tuyển ngay.",
        "Chọn Chia sẻ nếu muốn gửi tin cho người khác.",
        "Kiểm tra lại hồ sơ trước khi nộp.",
        "Sau khi nộp, theo dõi trạng thái ở khu vực hồ sơ cá nhân hoặc việc đã ứng tuyển.",
    ])

    doc.add_heading("5. Tìm kiếm công ty", level=1)
    doc.add_heading("5.1 Dùng trang công ty", level=2)
    add_paragraph(doc, "Trang Công ty cho phép bạn xem danh sách doanh nghiệp và tìm theo tên công ty.")
    add_bullets(doc, [
        "Tìm nhanh công ty phù hợp.",
        "Xem hồ sơ doanh nghiệp.",
        "Xem tin đang tuyển của công ty.",
    ])
    add_image(doc, "04-companies-top.png", "Hình 6. Trang danh sách công ty")

    doc.add_heading("5.2 Mở hồ sơ công ty", level=2)
    add_numbers(doc, [
        "Chọn một công ty trong danh sách.",
        "Đọc tên công ty, ngành hoạt động và năm thành lập.",
        "Xem website, email, số điện thoại và mô tả công ty.",
        "Xem các việc làm đang tuyển ngay tại hồ sơ công ty.",
    ])
    add_image(doc, "05-company-detail-top.png", "Hình 7. Trang chi tiết công ty")

    doc.add_heading("5.3 Theo dõi công ty", level=2)
    add_paragraph(doc, "Nếu bạn quan tâm một công ty, hãy mở hồ sơ và lưu lại để dễ quay lại sau.")

    doc.add_heading("6. Các thao tác thường dùng sau khi đăng nhập", level=1)
    doc.add_heading("6.1 Hồ sơ cá nhân", level=2)
    add_bullets(doc, [
        "Cập nhật ảnh đại diện.",
        "Bổ sung họ tên, số điện thoại và địa chỉ.",
        "Thêm kinh nghiệm, học vấn, kỹ năng và mục tiêu nghề nghiệp.",
        "Tải lên CV nếu có.",
    ])

    doc.add_heading("6.2 Việc đã lưu và việc đã ứng tuyển", level=2)
    add_bullets(doc, [
        "Kiểm tra lại các việc đã lưu để ứng tuyển sau.",
        "Xem trạng thái các việc đã ứng tuyển.",
        "Theo dõi tin nhắn hoặc thông báo từ hệ thống.",
    ])

    doc.add_heading("6.3 Thông báo", level=2)
    add_bullets(doc, [
        "Xác nhận đăng ký hoặc đăng nhập.",
        "Lời mời phỏng vấn.",
        "Cập nhật trạng thái hồ sơ.",
        "Thông báo từ nhà tuyển dụng hoặc hệ thống.",
    ])

    doc.add_heading("7. Phỏng vấn AI", level=1)
    add_paragraph(doc, "Một số vị trí có thể yêu cầu phỏng vấn AI. Quy trình thường như sau:")
    add_numbers(doc, [
        "Nhận lời mời phỏng vấn từ hệ thống.",
        "Kiểm tra micro, camera và kết nối Internet.",
        "Vào phòng phỏng vấn theo thời gian được chỉ định.",
        "Trả lời câu hỏi theo hướng dẫn trên màn hình hoặc qua giọng nói nếu được hỗ trợ.",
        "Kết thúc phỏng vấn và chờ hệ thống cập nhật kết quả.",
    ])
    add_bullets(doc, [
        "Ngồi ở nơi yên tĩnh trước khi bắt đầu.",
        "Chuẩn bị CV và thông tin ứng tuyển để trả lời thống nhất.",
        "Nếu bị ngắt kết nối, hãy vào lại bằng đường dẫn được cấp.",
    ])

    doc.add_heading("8. Dành cho nhà tuyển dụng", level=1)
    add_paragraph(doc, "Nhà tuyển dụng dùng các khu vực riêng sau khi đăng nhập vào cổng employer:")
    add_bullets(doc, [
        "Cập nhật hồ sơ công ty và thông tin liên hệ.",
        "Tạo, sửa và quản lý tin tuyển dụng.",
        "Tìm hồ sơ ứng viên và lưu ứng viên tiềm năng.",
        "Tạo nhóm câu hỏi, phỏng vấn AI và theo dõi kết quả.",
        "Xem lịch sử tương tác với ứng viên.",
    ])

    doc.add_heading("9. Xử lý lỗi cơ bản", level=1)
    doc.add_heading("9.1 Không đăng nhập được", level=2)
    add_bullets(doc, [
        "Kiểm tra lại email hoặc số điện thoại.",
        "Kiểm tra mật khẩu.",
        "Xác nhận tài khoản đã được kích hoạt chưa.",
        "Đặt lại mật khẩu nếu cần.",
    ])

    doc.add_heading("9.2 Không thấy dữ liệu", level=2)
    add_bullets(doc, [
        "Đổi từ khóa tìm kiếm.",
        "Bỏ bớt bộ lọc.",
        "Tải lại trang.",
        "Kiểm tra kết nối mạng.",
    ])

    doc.add_heading("9.3 Màn hình không phản hồi", level=2)
    add_bullets(doc, [
        "Làm mới trình duyệt.",
        "Đăng xuất và đăng nhập lại.",
        "Xóa cache nếu đã sử dụng lâu.",
        "Liên hệ hỗ trợ nếu lỗi lặp lại.",
    ])

    doc.add_heading("10. Kênh hỗ trợ", level=1)
    add_paragraph(doc, "Khi cần mở rộng tài liệu cho nội bộ, hãy bổ sung:")
    add_bullets(doc, [
        "Email hỗ trợ chính thức.",
        "Số hotline.",
        "Quy trình riêng cho ứng viên và nhà tuyển dụng.",
        "Ảnh chụp từng bước nếu muốn hướng dẫn chi tiết hơn cho người mới.",
    ])

    return doc


def main():
    if not IMG_DIR.exists():
        raise FileNotFoundError(f"Missing image directory: {IMG_DIR}")

    doc = build()
    doc.save(DOCX_PATH)
    print(f"Saved {DOCX_PATH}")


if __name__ == "__main__":
    main()

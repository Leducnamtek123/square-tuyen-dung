from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "HUONG_DAN_SU_DUNG.md"
DOCX_PATH = ROOT / "docs" / "HUONG_DAN_SU_DUNG.docx"


def set_document_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    if level:
        paragraph.paragraph_format.left_indent = Inches(0.25 * level)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str = "") -> None:
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Khu vực"
    hdr[1].text = "Đường dẫn"
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right


def build_document() -> Document:
    doc = Document()
    set_document_defaults(doc)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HƯỚNG DẪN SỬ DỤNG SQUARE TUYỂN DỤNG")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Phiên bản 1.0 - 09/05/2026")
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)

    doc.add_paragraph("")

    add_heading(doc, "1. Tổng quan hệ thống", 1)
    add_paragraph(
        doc,
        "Square Tuyển Dụng là nền tảng tuyển dụng gồm ba luồng chính: người tìm việc, nhà tuyển dụng và admin."
    )
    add_bullet(doc, "Người tìm việc xem tin, lưu tin, ứng tuyển và theo dõi lịch sử.")
    add_bullet(doc, "Nhà tuyển dụng đăng tin, quản lý hồ sơ ứng viên, tạo phỏng vấn AI và quản lý nội dung công ty.")
    add_bullet(doc, "Admin quản trị dữ liệu, kiểm duyệt nội dung và cấu hình hệ thống.")

    add_heading(doc, "Các cổng truy cập chính", 2)
    add_table(doc, [
        ("Trang công khai", "/"),
        ("Ứng viên", "/dashboard, /jobs, /companies, /profile, /my-jobs, /my-interviews"),
        ("Nhà tuyển dụng", "/employer/dashboard, /employer/job-posts, /employer/candidates, /employer/interviews"),
        ("Admin", "/admin/dashboard, /admin/users, /admin/jobs, /admin/interviews, /admin/settings"),
    ])

    add_heading(doc, "2. Hướng dẫn cho người dùng cuối", 1)
    add_heading(doc, "2.1 Đăng ký và đăng nhập", 2)
    for item in [
        "Mở trang chủ và chọn Đăng ký hoặc Đăng nhập.",
        "Chọn đúng loại tài khoản: Ứng viên hoặc Nhà tuyển dụng.",
        "Điền thông tin bắt buộc và xác nhận email nếu hệ thống yêu cầu.",
        "Sau khi đăng nhập, hệ thống đưa bạn đến trang tổng quan phù hợp.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2.2 Tạo và cập nhật hồ sơ cá nhân", 2)
    for item in [
        "Vào Hồ sơ hoặc Profile.",
        "Cập nhật thông tin cơ bản: họ tên, số điện thoại, ngày sinh, địa chỉ, kinh nghiệm, học vấn và kỹ năng.",
        "Tải lên CV hoặc tạo hồ sơ trực tuyến.",
        "Lưu lại để hệ thống dùng cho việc ứng tuyển.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2.3 Tìm việc", 2)
    for item in [
        "Vào Việc làm.",
        "Dùng bộ lọc theo từ khóa, ngành nghề, tỉnh thành và loại công việc.",
        "Mở chi tiết tin tuyển dụng để xem mô tả công việc, yêu cầu, quyền lợi và thông tin liên hệ.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2.4 Lưu tin và ứng tuyển", 2)
    for item in [
        "Mở một tin tuyển dụng phù hợp.",
        "Chọn Lưu tin nếu muốn xem lại sau.",
        "Chọn Ứng tuyển để nộp hồ sơ.",
        "Kiểm tra lại CV và thông tin liên hệ trước khi gửi.",
        "Theo dõi trạng thái ở Việc đã ứng tuyển hoặc My Jobs.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2.5 Theo dõi phỏng vấn AI", 2)
    for item in [
        "Khi được mời phỏng vấn, mở mục Phỏng vấn của tôi hoặc liên kết từ hệ thống.",
        "Kiểm tra thiết bị âm thanh và mạng trước khi vào phòng.",
        "Thực hiện phỏng vấn theo hướng dẫn trên màn hình.",
        "Sau khi kết thúc, xem lại trạng thái và lịch sử nếu hệ thống có hiển thị.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2.6 Nhắn tin, thông báo và công ty", 2)
    for item in [
        "Vào Thông báo để xem cập nhật mới.",
        "Vào Chat để trao đổi khi hệ thống cho phép.",
        "Vào Công ty để xem thông tin doanh nghiệp và bài viết liên quan.",
    ]:
        add_number(doc, item)

    add_heading(doc, "3. Hướng dẫn cho nhà tuyển dụng", 1)
    add_heading(doc, "3.1 Đăng ký và xác thực", 2)
    for item in [
        "Chọn Nhà tuyển dụng khi đăng ký.",
        "Điền thông tin doanh nghiệp và tài khoản.",
        "Hoàn tất xác minh nếu hệ thống yêu cầu.",
        "Sau khi vào cổng employer, hệ thống sẽ mở trang dashboard.",
    ]:
        add_number(doc, item)

    add_heading(doc, "3.2 Quản lý công ty", 2)
    for item in [
        "Vào Company hoặc Công ty.",
        "Cập nhật tên công ty, logo, ảnh bìa, mô tả, website và mạng xã hội.",
        "Quản lý thành viên nội bộ ở mục Employees nếu có quyền.",
    ]:
        add_number(doc, item)

    add_heading(doc, "3.3 Đăng và quản lý tin tuyển dụng", 2)
    for item in [
        "Vào Job Posts.",
        "Chọn Create để tạo tin mới.",
        "Nhập tên việc làm, số lượng, lương, địa điểm, mô tả, yêu cầu và quyền lợi.",
        "Lưu tin và chờ duyệt nếu quy trình yêu cầu.",
        "Theo dõi lượt xem, lượt chia sẻ và trạng thái tin.",
    ]:
        add_number(doc, item)

    add_heading(doc, "3.4 Tìm kiếm và quản lý ứng viên", 2)
    for item in [
        "Vào Candidates, Profiles hoặc Applied Profiles.",
        "Lọc theo kỹ năng, kinh nghiệm, vị trí hoặc hồ sơ đã ứng tuyển.",
        "Mở hồ sơ chi tiết để xem thông tin cá nhân, kinh nghiệm, học vấn, kỹ năng và chứng chỉ.",
        "Lưu hồ sơ ứng viên phù hợp để theo dõi sau.",
    ]:
        add_number(doc, item)

    add_heading(doc, "3.5 Tạo phỏng vấn AI", 2)
    for item in [
        "Vào Interviews.",
        "Chọn Create để tạo phiên phỏng vấn.",
        "Chọn bộ câu hỏi hoặc nhóm câu hỏi phù hợp.",
        "Thiết lập lịch hẹn, ứng viên và tin tuyển dụng.",
        "Chia sẻ link hoặc mã mời cho ứng viên.",
        "Theo dõi lịch sử, transcript và điểm đánh giá sau phỏng vấn.",
    ]:
        add_number(doc, item)

    add_heading(doc, "3.6 Quản lý nội dung và hỗ trợ", 2)
    for item in [
        "Vào Blog để tạo bài viết tuyển dụng hoặc thương hiệu nhà tuyển dụng.",
        "Vào Notifications để xem thông báo hệ thống.",
        "Vào Support nếu cần hỗ trợ.",
        "Vào Settings để cập nhật cấu hình tài khoản.",
    ]:
        add_number(doc, item)

    add_heading(doc, "4. Hướng dẫn cho admin", 1)
    add_heading(doc, "4.1 Đăng nhập", 2)
    for item in [
        "Truy cập /admin/login.",
        "Đăng nhập bằng tài khoản có quyền quản trị.",
        "Sau khi vào hệ thống, mở Dashboard để xem tổng quan.",
    ]:
        add_number(doc, item)

    add_heading(doc, "4.2 Các khu vực quản trị chính", 2)
    for item in [
        "Users",
        "Jobs",
        "Profiles",
        "Companies",
        "Resumes",
        "Interviews",
        "Question Bank",
        "Question Groups",
        "Articles",
        "Banners",
        "Banner Types",
        "Feedbacks",
        "Cities",
        "Districts",
        "Wards",
        "Job Activity",
        "Job Notifications",
        "Settings",
        "Chat",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.3 Quản lý người dùng", 2)
    for item in [
        "Vào Users.",
        "Kiểm tra trạng thái tài khoản: hoạt động, xác minh email, quyền staff hoặc superuser.",
        "Cập nhật thông tin, đổi mật khẩu hoặc bật/tắt tài khoản khi cần.",
    ]:
        add_number(doc, item)

    add_heading(doc, "4.4 Quản lý tin tuyển dụng", 2)
    for item in [
        "Vào Jobs.",
        "Duyệt nội dung tin và kiểm tra tính hợp lệ.",
        "Thay đổi trạng thái tin khi cần: chờ duyệt, đã duyệt hoặc từ chối.",
        "Theo dõi lượt xem, lượt chia sẻ và lịch sử thay đổi.",
    ]:
        add_number(doc, item)

    add_heading(doc, "4.5 Quản lý phỏng vấn", 2)
    for item in [
        "Vào Interviews.",
        "Kiểm tra danh sách phiên phỏng vấn, transcript và đánh giá.",
        "Mở Question Bank và Question Groups để quản lý bộ câu hỏi.",
        "Kiểm tra thống kê phỏng vấn nếu cần.",
    ]:
        add_number(doc, item)

    add_heading(doc, "4.6 Quản lý nội dung và giao diện", 2)
    for item in [
        "Vào Articles để tạo hoặc sửa bài viết.",
        "Vào Banners và Banner Types để quản lý banner hiển thị.",
        "Vào Feedbacks để duyệt phản hồi người dùng.",
        "Vào Settings để cấu hình các tham số hệ thống.",
    ]:
        add_number(doc, item)

    add_heading(doc, "4.7 Quản lý địa danh", 2)
    for item in [
        "Vào Cities, Districts và Wards.",
        "Kiểm tra dữ liệu địa lý dùng cho bộ lọc và hồ sơ.",
        "Cập nhật khi có thay đổi dữ liệu hành chính.",
    ]:
        add_number(doc, item)

    add_heading(doc, "5. Lưu ý vận hành", 1)
    for item in [
        "Tài khoản phải đúng loại quyền mới thấy đúng portal.",
        "Một số chức năng có thể yêu cầu email đã xác minh.",
        "Khi upload ảnh hoặc CV, nên dùng định dạng phổ biến như JPG, PNG, PDF.",
        "Nếu phỏng vấn AI không vào được, kiểm tra lại micro, camera và kết nối mạng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Hỗ trợ", 1)
    add_paragraph(doc, "Nếu cần chỉnh sửa thêm theo đúng quy trình nội bộ, nên bổ sung:")
    for item in [
        "Tên thương hiệu chính thức",
        "Ảnh chụp màn hình từng bước",
        "Số hotline hoặc email hỗ trợ",
        "Quy trình phê duyệt nội bộ của doanh nghiệp",
    ]:
        add_bullet(doc, item)

    return doc


def main() -> None:
    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)
    doc = build_document()
    doc.save(DOCX_PATH)
    print(f"Saved {DOCX_PATH}")


if __name__ == "__main__":
    main()
